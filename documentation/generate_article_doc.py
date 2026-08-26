"""Generate a Word document for the CloudFront Cognito proxy article."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- Title ---
title = doc.add_heading(
    "Fixing AWS Cognito SChannel TLS Failures on Windows Using CloudFront as a Reverse Proxy",
    level=0,
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- Introduction ---
doc.add_heading("The Problem", level=1)

doc.add_paragraph(
    "Single-page applications (SPAs) that authenticate users with Amazon Cognito's "
    "Hosted UI must connect to AWS-managed regional endpoints. On Windows, the operating "
    "system's native TLS implementation — SChannel — handles these connections. In certain "
    "environments, SChannel fails to complete the TLS handshake with Cognito's endpoints, "
    "resulting in connection errors that prevent users from logging in."
)

doc.add_paragraph(
    "This is particularly frustrating because the Cognito endpoint's TLS configuration "
    "is entirely managed by AWS. You cannot change the cipher suites, certificate chain, "
    "or TLS version offered by the Cognito service. If SChannel on the client machine "
    "cannot negotiate a session, the user simply cannot authenticate."
)

# --- Symptoms ---
doc.add_heading("What Users See", level=1)

doc.add_paragraph(
    "The symptoms vary depending on how the SPA handles the failure, but users typically "
    "report one or more of the following:"
)

bullets = [
    'The browser shows "This site can\'t be reached" or "Connection was reset" when '
    "redirected to the Cognito login page.",
    "A blank page or infinite spinner after clicking the Login button, because the "
    "OAuth2 authorize redirect never completes.",
    "An ERR_SSL_PROTOCOL_ERROR or ERR_CONNECTION_RESET error in Chrome/Edge on Windows, "
    "while the same flow works on macOS or Linux.",
    "The token exchange (POST to /oauth2/token) fails silently — the authorize redirect "
    "may succeed, but the background fetch() call to exchange the authorization code "
    "for tokens is rejected by SChannel before it reaches Cognito.",
    "Intermittent failures that depend on Windows version, patch level, or Group Policy "
    "TLS settings applied by enterprise IT.",
]

for bullet in bullets:
    doc.add_paragraph(bullet, style="List Bullet")

doc.add_paragraph(
    "Critically, the same user on the same machine can often reach other HTTPS sites "
    "without issue. The problem is specific to the TLS configuration on the Cognito "
    "endpoint, which may use cipher suites or certificate chains that certain SChannel "
    "configurations do not support."
)

# --- Root Cause ---
doc.add_heading("Root Cause", level=1)

doc.add_paragraph(
    "SChannel is the TLS/SSL provider built into Windows. Unlike OpenSSL (used by most "
    "Linux/macOS applications), SChannel's behavior is controlled by the Windows registry, "
    "Group Policy, and Windows Update patches. In enterprise environments, IT departments "
    "often restrict which TLS versions and cipher suites SChannel will negotiate."
)

doc.add_paragraph(
    "AWS Cognito's regional endpoints (e.g., *.auth.<region>.amazoncognito.com) present "
    "a TLS configuration that is optimized for broad compatibility but is not configurable "
    "by the customer. When the cipher suites or TLS parameters offered by Cognito do not "
    "intersect with what a particular Windows SChannel configuration allows, the handshake "
    "fails at the OS level — before any HTTP traffic is exchanged."
)

doc.add_paragraph(
    "Because the Cognito endpoint is fully AWS-managed, there is no customer-side fix "
    "on the server. The options are:"
)

options = [
    "Modify the Windows client's SChannel configuration (often not possible in enterprise "
    "environments due to Group Policy).",
    "Route traffic through an intermediary that terminates TLS with a configuration you "
    "control — such as Amazon CloudFront.",
]

for opt in options:
    doc.add_paragraph(opt, style="List Number")

# --- Solution ---
doc.add_heading("The Fix: CloudFront as a Reverse Proxy for Cognito", level=1)

doc.add_paragraph(
    "If your SPA is already served from a CloudFront distribution (common for S3-hosted "
    "SPAs), you can add Cognito as an additional origin behind a specific path pattern. "
    "The SPA then sends its authentication requests to your own CloudFront domain instead "
    "of directly to Cognito. CloudFront proxies those requests to Cognito on the backend."
)

doc.add_paragraph("This works because:")

reasons = [
    "CloudFront's TLS configuration for viewer connections is customer-configurable — "
    "you choose the security policy, which determines cipher suites and TLS versions.",
    "The connection between CloudFront and the Cognito origin (backend) is managed by "
    "AWS infrastructure on the AWS network, which does not have the SChannel limitation.",
    "From the user's perspective, all traffic goes to your domain "
    "(e.g., app.yourdomain.com), eliminating the problematic direct connection to "
    "the Cognito endpoint.",
]

for reason in reasons:
    doc.add_paragraph(reason, style="List Bullet")

# --- Architecture ---
doc.add_heading("Architecture", level=1)

doc.add_paragraph("Before (direct connection):")
doc.add_paragraph(
    "    Browser (SChannel) → cognito-idp.<region>.amazoncognito.com  [TLS FAILURE]",
    style="No Spacing",
)

doc.add_paragraph("")
doc.add_paragraph("After (proxied through CloudFront):")
doc.add_paragraph(
    "    Browser (SChannel) → app.yourdomain.com/cognito/*  [CloudFront, TLS OK]",
    style="No Spacing",
)
doc.add_paragraph(
    "    CloudFront → *.auth.<region>.amazoncognito.com     [AWS backbone, TLS OK]",
    style="No Spacing",
)

# --- Implementation ---
doc.add_heading("Implementation Steps", level=1)

doc.add_heading("1. Identify Your Cognito Domain", level=2)

doc.add_paragraph(
    "Look up the exact Cognito Hosted UI domain from your CloudFormation stack outputs "
    "(key: CognitoDomain) or from the Cognito console under your user pool's domain "
    "settings. The hostname follows the pattern:"
)
doc.add_paragraph(
    "    <domain-prefix>.auth.<region>.amazoncognito.com", style="No Spacing"
)
doc.add_paragraph(
    "Verify the domain resolves in DNS before proceeding. An NXDOMAIN response means "
    "the prefix or region is incorrect."
)

doc.add_heading("2. Add a CloudFront Origin", level=2)

doc.add_paragraph(
    "In your CloudFront distribution, create a new origin pointing to the Cognito "
    "Hosted UI hostname. Use HTTPS-only, port 443, minimum TLS 1.2. Leave the origin "
    "path empty."
)

doc.add_heading("3. Create a Cache Behavior", level=2)

doc.add_paragraph("Add a behavior with these settings:")

behavior_items = [
    "Path pattern: /cognito/*",
    "Origin: the Cognito origin created above",
    "Allowed HTTP methods: ALL (GET, HEAD, OPTIONS, PUT, POST, PATCH, DELETE)",
    "Cache policy: CachingDisabled (Cognito requests are user-specific)",
    'Origin request policy: AllViewerExceptHostHeader — this is critical. CloudFront '
    "must send Cognito's own hostname as the Host header, not your distribution's domain. "
    "This managed policy forwards all viewer headers except Host, letting CloudFront "
    "substitute the correct origin hostname.",
    "Viewer protocol policy: HTTPS only",
]

for item in behavior_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4. Add a CloudFront Function for Path Rewriting", level=2)

doc.add_paragraph(
    "Your SPA sends requests to /cognito/oauth2/token, but Cognito expects /oauth2/token. "
    "Create a CloudFront Function (viewer-request stage) that strips the /cognito prefix:"
)

doc.add_paragraph(
    '    function handler(event) {\n'
    '        var request = event.request;\n'
    "        request.uri = request.uri.replace(/^\\/cognito/, '') || '/';\n"
    '        return request;\n'
    '    }',
    style="No Spacing",
)

doc.add_paragraph(
    "Publish the function and associate it with the /cognito/* behavior as a viewer-request "
    "function."
)

doc.add_heading("5. Update the SPA Configuration", level=2)

doc.add_paragraph(
    "Change the Cognito domain URL in your frontend configuration from the direct "
    "Cognito endpoint to the CloudFront proxy path:"
)

doc.add_paragraph(
    "    Before:  https://<prefix>.auth.<region>.amazoncognito.com", style="No Spacing"
)
doc.add_paragraph(
    "    After:   https://app.yourdomain.com/cognito", style="No Spacing"
)

doc.add_paragraph(
    "If your SPA's auth library constructs OAuth2 URLs by appending paths like "
    "/oauth2/authorize and /oauth2/token to a base domain, this single configuration "
    "change routes all auth traffic through CloudFront automatically."
)

# --- Key Technical Details ---
doc.add_heading("Key Technical Details", level=1)

doc.add_heading("Why AllViewerExceptHostHeader is Required", level=2)

doc.add_paragraph(
    "Cognito validates the Host header on incoming requests. If CloudFront forwards "
    "Host: app.yourdomain.com (the viewer's Host header), Cognito rejects the request "
    "because it does not recognize that hostname. The AllViewerExceptHostHeader origin "
    "request policy strips the viewer's Host header and lets CloudFront send the "
    "origin's hostname instead, which Cognito accepts."
)

doc.add_paragraph(
    "Using the wrong policy (e.g., AllViewer) results in 502 or 403 errors from "
    "Cognito, even though CloudFront successfully connects to the origin."
)

doc.add_heading("Why CachingDisabled is Required", level=2)

doc.add_paragraph(
    "Every Cognito OAuth2 request contains user-specific parameters (authorization codes, "
    "PKCE challenges, tokens). Caching any of these responses would break authentication "
    "or leak tokens between users."
)

doc.add_heading("Hosted UI Login Page Behavior", level=2)

doc.add_paragraph(
    "The /oauth2/authorize endpoint returns a 302 redirect to Cognito's login page. "
    "This redirect's Location header points to the direct Cognito domain — the browser "
    "will follow it and load the Hosted UI directly from Cognito. This means the initial "
    "login page may still connect directly to Cognito."
)

doc.add_paragraph(
    "However, the token exchange (/oauth2/token) — which is where SChannel issues most "
    "commonly manifest as a fetch() failure — goes through the proxy. If the SChannel "
    "failure also affects the Hosted UI page load, you would need to replace the Hosted "
    "UI with a custom login form that calls the Cognito API through the proxy."
)

# --- Troubleshooting ---
doc.add_heading("Troubleshooting", level=1)

trouble_items = [
    ("503 from CloudFront", "The CloudFront Function is not published or not associated "
     "with the behavior. Verify it is published (not draft) and attached as a "
     "viewer-request function."),
    ("502 from CloudFront", "TLS handshake failure between CloudFront and the Cognito "
     "origin. Usually means the origin domain is wrong (NXDOMAIN) or the Host header "
     "is being forwarded incorrectly. Check the origin domain resolves and the origin "
     "request policy is AllViewerExceptHostHeader."),
    ("403 from Cognito", "Cognito is receiving an unexpected Host header or the request "
     "path is wrong. Verify the CloudFront Function is stripping the prefix correctly."),
    ("NXDOMAIN on the Cognito domain", "The domain prefix or region is incorrect. Look "
     "up the actual domain from CloudFormation stack outputs."),
]

for title_text, desc in trouble_items:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(title_text + ": ")
    run.bold = True
    p.add_run(desc)

# --- Conclusion ---
doc.add_heading("Conclusion", level=1)

doc.add_paragraph(
    "By leveraging CloudFront as a reverse proxy for Cognito, you eliminate the direct "
    "TLS dependency between Windows SChannel and the AWS-managed Cognito endpoint. The "
    "fix requires no changes to Cognito itself, no modification of Windows TLS settings, "
    "and no custom authentication backend. It works entirely within the AWS edge "
    "infrastructure you already use to serve your SPA."
)

# --- Save ---
output_path = r"c:\Users\mgeis\Downloads\Agentic-compare\Kiro\kiro-agentcore-app\documentation\cloudfront-cognito-proxy-article.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
